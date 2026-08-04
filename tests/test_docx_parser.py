from __future__ import annotations

from pathlib import Path

from docx import Document

from evenor.adapters.rag.parsers.registry import parse_file, supported_suffixes


def test_parse_docx_roundtrip(tmp_path: Path) -> None:
    doc_path = tmp_path / "note.docx"
    doc = Document()
    doc.add_paragraph("Title line")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "ColA"
    table.cell(0, 1).text = "ColB"
    doc.add_paragraph("After table")
    doc.save(doc_path)

    text = parse_file(doc_path)
    assert "Title line" in text
    assert "ColA | ColB" in text
    assert "After table" in text


def test_docx_in_supported_suffixes() -> None:
    assert ".docx" in supported_suffixes()
